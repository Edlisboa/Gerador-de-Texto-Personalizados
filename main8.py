import re
import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT
from openpyxl import load_workbook

def formatar_milhar(numero):
    if numero is not None:
        try:
            return "{:,}".format(int(numero)).replace(",", ".")
        except ValueError:
            return str(numero)
    return ""

def formatar_data(data):
    if data:
        try:
            from datetime import datetime
            data_formatada = datetime.strptime(str(data), "%Y-%m-%d %H:%M:%S")
            return data_formatada.strftime("%d/%m/%Y")
        except ValueError:
            return str(data)
    return ""

def preencher_texto(texto_base, dados):
    texto_preenchido = texto_base
    for variavel, valor in dados.items():
        texto_preenchido = texto_preenchido.replace(f"[{variavel}]", str(valor))
    return texto_preenchido

# Texto base fornecido
texto_base = """AV-**/[mat] - CANCELAMENTO DE INDISPONIBILIDADE - Em [data_selagem]. Protocolado sob o nº [n_protocolo], em [data_protocolo]. Nos termos do protocolo de cancelamento do CNIB nº [prot_cnib], datado de [data_cnib], expedido pelo Tribunal Superior do Trabalho, Tribunal Regional do Trabalho da 23ª Região, [n_vara]ª Vara do Trabalho de Cuiabá-MT, referente aos autos do processo nº [n_processo], da vara supracitada, procede-se a presente averbação para baixar a restrição de indisponibilidade constante da AV-** acima. SELO DE AUTENTICIDADE: [selo_digital]. Não foram cobrados emolumentos por força da Lei. Eu,__________Oficial que fiz digitar e conferi.
____________________________________________________________________________"""

# Caminho do arquivo Excel
desktop = os.path.join(os.path.join(os.environ['USERPROFILE']), 'Desktop')
caminho_excel = os.path.join(desktop, 'minuta automatica', 'produção', 'dados.xlsx')

# Carrega o arquivo Excel e seleciona a "Planilha 1"
workbook = load_workbook(caminho_excel)
sheet = workbook["Planilha1"]

# Cria um novo documento Word
document = Document()

# Aplica formatação de página
section = document.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = Cm(2.0)
section.right_margin = Cm(1.2)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.orientation = WD_ORIENT.PORTRAIT

# Aplica formatação de fonte padrão
estilo_paragrafo = document.styles['Normal'].paragraph_format
estilo_fonte = document.styles['Normal'].font
estilo_fonte.name = 'Courier New'
estilo_fonte.size = Pt(11)

# Itera sobre as linhas do Excel
for row in sheet.iter_rows(min_row=2, values_only=True):
    dados = {
        "mat": formatar_milhar(row[0]),
        "data_selagem": formatar_data(row[1]),
        "n_protocolo": row[2],
        "data_protocolo": formatar_data(row[3]),
        "prot_cnib": row[4],
        "data_cnib": formatar_data(row[5]),
        "n_vara": row[6],
        "n_processo": row[7],
        "selo_digital": row[8],
    }
    # Verifica se o valor de selo_digital é válido
    if dados['selo_digital']:
        texto_preenchido = preencher_texto(texto_base, dados)

        # Define as posições dos trechos em negrito
        posicoes_negrito = [
            (0, texto_preenchido.find(" - Em ")),
            (texto_preenchido.find("baixar a restrição de indisponibilidade"), texto_preenchido.find(" constante da AV-** acima.")),
            (texto_preenchido.find("AV-** acima."), texto_preenchido.find("SELO DE AUTENTICIDADE: ")),
            (texto_preenchido.find(dados['selo_digital']), texto_preenchido.find(". Não foram cobrados emolumentos")),
            (texto_preenchido.find("____________________________________________________________________________"), len(texto_preenchido))
        ]

        # Adiciona os parágrafos com formatação em negrito
        p = document.add_paragraph()
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        p.paragraph_format.left_indent = Cm(0)
        p.paragraph_format.right_indent = Cm(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0

        pos_atual = 0
        for inicio, fim in posicoes_negrito:
            if inicio > pos_atual:
                p.add_run(texto_preenchido[pos_atual:inicio])
            run = p.add_run(texto_preenchido[inicio:fim])
            run.bold = True
            pos_atual = fim
        if pos_atual < len(texto_preenchido):
            p.add_run(texto_preenchido[pos_atual:])
    
        # Adiciona duas quebras de linha para separar as minutas
        if row != list(sheet.iter_rows(min_row=2, values_only=True))[-1]:
            document.add_paragraph()
            document.add_paragraph()
    else:
        print(f"Aviso: selo_digital vazio na linha {row}")

# Define o caminho para salvar o arquivo .docx na pasta "resultados"
caminho_pasta_resultados = os.path.join(desktop, 'minuta automatica', 'resultados')
os.makedirs(caminho_pasta_resultados, exist_ok=True)  # Cria a pasta se não existir

caminho_arquivo_base = os.path.join(caminho_pasta_resultados, 'minutas.docx')

# Verifica se o arquivo já existe e gera um novo nome se necessário
if os.path.exists(caminho_arquivo_base):
    contador = 1
    while True:
        caminho_arquivo = os.path.splitext(caminho_arquivo_base)[0] + f"_{contador}.docx"
        if not os.path.exists(caminho_arquivo):
            break
        contador += 1
else:
    caminho_arquivo = caminho_arquivo_base

# Salva o documento
document.save(caminho_arquivo)
print(f"Arquivo '{os.path.basename(caminho_arquivo)}' salvo em: {caminho_arquivo}")