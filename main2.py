import re
import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT
from openpyxl import load_workbook
from datetime import datetime

def formatar_data(data, variavel):
    if isinstance(data, datetime) and variavel in ["data_selagem", "data_protocolo"]:
        return data.strftime("%d/%m/%Y")
    return str(data)

def preencher_texto(texto_base, dados):
    texto_preenchido = texto_base
    for variavel, valor in dados.items():
        texto_preenchido = texto_preenchido.replace(f"[{variavel}]", str(valor))
    return texto_preenchido

# Texto base fornecido
texto_base = """AV-**/[mat] - CANCELAMENTO DE INDISPONIBILIDADE - Em [data_selagem]. Protocolado sob o nº [n_protocolo], em [data_protocolo]. Nos termos do protocolo de cancelamento do CNIB nº [prot_cnib], datado de [data_cnib], expedido pelo Tribunal Superior do Trabalho, Tribunal Regional do Trabalho da 23ª Região, [n_vara]ª Vara do Trabalho de Cuiabá-MT, referente aos autos do processo nº [n_processo], da vara supracitada, procede-se a presente averbação para baixar a restrição de indisponibilidade constante da AV-** acima. SELO DE AUTENTICIDADE: [selo_digital]. Não foram cobrados emolumentos por força da Lei. Eu,__________Oficial que fiz digitar e conferi.
____________________________________________________________________________"""

# Caminho do arquivo Excel
desktop = os.path.join(os.path.join(os.environ['USERPROFILE']), 'Desktop')
caminho_excel = os.path.join(desktop, 'minuta automatica', 'dados.xlsx')

# Carrega o arquivo Excel
workbook = load_workbook(caminho_excel)
sheet = workbook.active

# Cria um novo documento Word
document = Document()

# Aplica formatação de página
section = document.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = Cm(2.0)
section.right_margin = Cm(1.2)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.orientation = WD_ORIENT.PORTRAIT

# Aplica formatação de fonte padrão
estilo_paragrafo = document.styles['Normal'].paragraph_format
estilo_fonte = document.styles['Normal'].font
estilo_fonte.name = 'Courier New'
estilo_fonte.size = Pt(11)

# Itera sobre as linhas do Excel
for row in sheet.iter_rows(min_row=2, values_only=True):
    dados = {
        "mat": row[0],
        "data_selagem": formatar_data(row[1], "data_selagem"),
        "n_protocolo": row[2],
        "data_protocolo": formatar_data(row[3], "data_protocolo"),
        "prot_cnib": row[4],
        "data_cnib": formatar_data(row[5], "data_cnib"),
        "n_vara": row[6],
        "n_processo": row[7],
        "selo_digital": row[8],
    }

    texto_preenchido = preencher_texto(texto_base, dados)

    # Separa as partes do texto para aplicar o negrito
    partes = texto_preenchido.split("baixar a restrição de indisponibilidade")
    parte1 = partes[0]
    parte2 = partes[1]

    partes_av_acima = parte2.split("AV-** acima.")
    parte2_1 = partes_av_acima[0]
    parte2_2 = partes_av_acima[1]

    partes_selo = parte2_2.split("SELO DE AUTENTICIDADE: ")
    parte2_2_1 = partes_selo[0]
    parte2_2_2 = partes_selo[1]

    partes_selo_input = parte2_2_2.split('.')
    selo_input = partes_selo_input[0]
    selo_resto = '.' + partes_selo_input[1] if len(partes_selo_input)>1 else ''

    # Adiciona os parágrafos com formatação em negrito
    p = document.add_paragraph()
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p.paragraph_format.left_indent = Cm(0)
    p.paragraph_format.right_indent = Cm(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0

    run1 = p.add_run(parte1)
    run1.bold = True

    run2 = p.add_run("baixar a restrição de indisponibilidade")
    run2.bold = True

    run3 = p.add_run(parte2_1)

    run4 = p.add_run("AV-** acima.")
    run4.bold = True

    run5 = p.add_run(parte2_2_1)

    run6 = p.add_run("SELO DE AUTENTICIDADE: " + selo_input)
    run6.bold = True

    run7 = p.add_run(selo_resto)

    # Adiciona a parte faltante do texto com o ponto final
    p.add_run(".\nEu,__________Oficial que fiz digitar e conferi.")
    p.add_run("\n") # Adiciona uma quebra de linha após o ponto final

    # Adiciona os underlines
    p.add_run("\n")
    p.add_run("____________________________________________________________________________")

    # Adiciona duas quebras de linha para separar as minutas
    if row != list(sheet.iter_rows(min_row=2, values_only=True))[-1]:
        document.add_paragraph()
        document.add_paragraph()

# Define o caminho para salvar o arquivo .docx na área de trabalho
caminho_arquivo_base = os.path.join(desktop, 'minuta automatica', 'minutas.docx')

# Verifica se o arquivo já existe e gera um novo nome se necessário
if os.path.exists(caminho_arquivo_base):
    contador = 1
    while True:
        caminho_arquivo = os.path.splitext(caminho_arquivo_base)[0] + f"_{contador}.docx"
        if not os.path.exists(caminho_arquivo):
            break
        contador += 1
else:
    caminho_arquivo = caminho_arquivo_base

# Cria a pasta 'minuta automatica' se ela não existir
os.makedirs(os.path.dirname(caminho_arquivo), exist_ok=True)

# Salva o documento
document.save(caminho_arquivo)

print(f"Arquivo '{os.path.basename(caminho_arquivo)}' salvo em: {caminho_arquivo}")